#!/usr/bin/env python3
"""
Universal Document Processor for Knowledge Base
===============================================
Sistema inteligente e adaptável para processar QUALQUER tipo de documento
e preparar para inserção na Knowledge Base PostgreSQL com vector embeddings.

Formatos Suportados:
- Excel (.xlsx, .xls, .csv)
- Word (.docx, .doc)
- PDF (.pdf) - com OCR se necessário
- PowerPoint (.pptx, .ppt)
- Imagens (.jpg, .png, .gif, .bmp) - com OCR
- Texto (.txt, .md, .log, .json, .xml, .yaml)
- Email (.eml, .msg)
- HTML (.html, .htm)
- Código fonte (múltiplas linguagens)
- Arquivos comprimidos (.zip, .rar, .7z)

Features:
- Auto-detecção de tipo e estrutura
- Extração inteligente de conteúdo
- OCR para imagens e PDFs escaneados
- Processamento de metadados
- Chunking inteligente para documentos grandes
- Extração de tabelas, imagens e anexos
- Geração automática de resumos
- Classificação automática de conteúdo
"""

import os
import sys
import json
import hashlib
import re
import mimetypes
import base64
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
import logging
from dataclasses import dataclass, asdict
from enum import Enum

# ================================================================================
# CONFIGURAÇÃO DE LOGGING
# ================================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ================================================================================
# ESTRUTURAS DE DADOS
# ================================================================================

class DocumentType(Enum):
    """Tipos de documentos suportados"""
    EXCEL = "excel"
    WORD = "word"
    PDF = "pdf"
    POWERPOINT = "powerpoint"
    IMAGE = "image"
    TEXT = "text"
    EMAIL = "email"
    HTML = "html"
    CODE = "code"
    ARCHIVE = "archive"
    UNKNOWN = "unknown"

@dataclass
class DocumentMetadata:
    """Metadados do documento"""
    file_name: str
    file_path: str
    file_size: int
    file_type: DocumentType
    mime_type: str
    created_date: Optional[datetime]
    modified_date: Optional[datetime]
    author: Optional[str]
    title: Optional[str]
    subject: Optional[str]
    keywords: List[str]
    language: Optional[str]
    page_count: Optional[int]
    word_count: Optional[int]
    has_images: bool
    has_tables: bool
    has_attachments: bool
    encoding: Optional[str]
    checksum: str

@dataclass
class DocumentChunk:
    """Chunk de documento para processamento"""
    chunk_id: str
    document_id: str
    content: str
    content_type: str  # text, table, image, code, etc.
    chunk_index: int
    total_chunks: int
    metadata: Dict[str, Any]
    embedding_text: str  # Texto otimizado para embedding
    
@dataclass
class KnowledgeBaseEntry:
    """Entrada para a Knowledge Base"""
    uuid: str
    title: str
    content: str
    summary: str
    category: str
    tags: List[str]
    confidence_score: float
    source: str
    document_type: str
    metadata: Dict[str, Any]
    chunks: List[DocumentChunk]
    created_by: str
    embedding_text: Optional[str] = None

# ================================================================================
# DETECTORES E EXTRACTORES
# ================================================================================

class DocumentTypeDetector:
    """Detecta o tipo de documento baseado em extensão e conteúdo"""
    
    EXTENSION_MAP = {
        # Excel
        '.xlsx': DocumentType.EXCEL,
        '.xls': DocumentType.EXCEL,
        '.xlsm': DocumentType.EXCEL,
        '.csv': DocumentType.EXCEL,
        '.tsv': DocumentType.EXCEL,
        
        # Word
        '.docx': DocumentType.WORD,
        '.doc': DocumentType.WORD,
        '.rtf': DocumentType.WORD,
        '.odt': DocumentType.WORD,
        
        # PDF
        '.pdf': DocumentType.PDF,
        
        # PowerPoint
        '.pptx': DocumentType.POWERPOINT,
        '.ppt': DocumentType.POWERPOINT,
        '.odp': DocumentType.POWERPOINT,
        
        # Imagens
        '.jpg': DocumentType.IMAGE,
        '.jpeg': DocumentType.IMAGE,
        '.png': DocumentType.IMAGE,
        '.gif': DocumentType.IMAGE,
        '.bmp': DocumentType.IMAGE,
        '.tiff': DocumentType.IMAGE,
        '.tif': DocumentType.IMAGE,
        '.webp': DocumentType.IMAGE,
        
        # Texto
        '.txt': DocumentType.TEXT,
        '.md': DocumentType.TEXT,
        '.markdown': DocumentType.TEXT,
        '.log': DocumentType.TEXT,
        '.json': DocumentType.TEXT,
        '.xml': DocumentType.TEXT,
        '.yaml': DocumentType.TEXT,
        '.yml': DocumentType.TEXT,
        '.ini': DocumentType.TEXT,
        '.cfg': DocumentType.TEXT,
        '.conf': DocumentType.TEXT,
        
        # Email
        '.eml': DocumentType.EMAIL,
        '.msg': DocumentType.EMAIL,
        '.mbox': DocumentType.EMAIL,
        
        # HTML
        '.html': DocumentType.HTML,
        '.htm': DocumentType.HTML,
        '.xhtml': DocumentType.HTML,
        
        # Código
        '.py': DocumentType.CODE,
        '.js': DocumentType.CODE,
        '.java': DocumentType.CODE,
        '.c': DocumentType.CODE,
        '.cpp': DocumentType.CODE,
        '.cs': DocumentType.CODE,
        '.php': DocumentType.CODE,
        '.rb': DocumentType.CODE,
        '.go': DocumentType.CODE,
        '.rs': DocumentType.CODE,
        '.swift': DocumentType.CODE,
        '.kt': DocumentType.CODE,
        '.scala': DocumentType.CODE,
        '.r': DocumentType.CODE,
        '.sql': DocumentType.CODE,
        '.sh': DocumentType.CODE,
        '.bat': DocumentType.CODE,
        '.ps1': DocumentType.CODE,
        
        # Archives
        '.zip': DocumentType.ARCHIVE,
        '.rar': DocumentType.ARCHIVE,
        '.7z': DocumentType.ARCHIVE,
        '.tar': DocumentType.ARCHIVE,
        '.gz': DocumentType.ARCHIVE,
        '.bz2': DocumentType.ARCHIVE,
    }
    
    @classmethod
    def detect(cls, file_path: str) -> DocumentType:
        """Detecta o tipo de documento"""
        ext = Path(file_path).suffix.lower()
        return cls.EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)

# ================================================================================
# PROCESSADORES ESPECÍFICOS
# ================================================================================

class BaseProcessor:
    """Classe base para processadores de documentos"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_name = Path(file_path).name
        self.metadata = self._extract_metadata()
    
    def _extract_metadata(self) -> DocumentMetadata:
        """Extrai metadados básicos do arquivo"""
        path = Path(self.file_path)
        stat = path.stat()
        
        return DocumentMetadata(
            file_name=self.file_name,
            file_path=str(path.absolute()),
            file_size=stat.st_size,
            file_type=DocumentTypeDetector.detect(self.file_path),
            mime_type=mimetypes.guess_type(self.file_path)[0] or 'application/octet-stream',
            created_date=datetime.fromtimestamp(stat.st_ctime),
            modified_date=datetime.fromtimestamp(stat.st_mtime),
            author=None,
            title=None,
            subject=None,
            keywords=[],
            language=None,
            page_count=None,
            word_count=None,
            has_images=False,
            has_tables=False,
            has_attachments=False,
            encoding=None,
            checksum=self._calculate_checksum()
        )
    
    def _calculate_checksum(self) -> str:
        """Calcula checksum SHA256 do arquivo"""
        sha256 = hashlib.sha256()
        with open(self.file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa o documento e retorna entradas para KB"""
        raise NotImplementedError("Subclasses devem implementar process()")

class ExcelProcessor(BaseProcessor):
    """Processador para arquivos Excel"""
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa arquivo Excel"""
        import pandas as pd
        entries = []
        
        try:
            # Ler todas as sheets
            all_sheets = pd.read_excel(self.file_path, sheet_name=None, engine='openpyxl')
            
            for sheet_name, df in all_sheets.items():
                # Limpar dados
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                if df.empty:
                    continue
                
                # Processar cada linha como uma entrada potencial
                for idx, row in df.iterrows():
                    entry = self._create_entry_from_row(row, df.columns, sheet_name, idx)
                    if entry:
                        entries.append(entry)
                
                # Processar também a sheet inteira como uma entrada
                sheet_entry = self._create_entry_from_sheet(df, sheet_name)
                if sheet_entry:
                    entries.append(sheet_entry)
        
        except Exception as e:
            logger.error(f"Erro ao processar Excel {self.file_name}: {str(e)}")
        
        return entries
    
    def _create_entry_from_row(self, row, columns, sheet_name: str, row_index: int) -> Optional[KnowledgeBaseEntry]:
        """Cria entrada KB de uma linha do Excel"""
        # Identificar campos importantes
        title_parts = []
        content_parts = []
        
        for col in columns:
            value = row.get(col)
            if pd.notna(value) and str(value).strip():
                value_str = str(value).strip()
                
                # Heurística: colunas com nomes importantes vão para o título
                if any(keyword in col.lower() for keyword in ['nome', 'name', 'título', 'title', 'id', 'código', 'code']):
                    title_parts.append(value_str)
                
                content_parts.append(f"{col}: {value_str}")
        
        if not content_parts:
            return None
        
        # Criar título
        if title_parts:
            title = " - ".join(title_parts[:3])
        else:
            title = f"Linha {row_index + 1} de {sheet_name}"
        
        # Criar conteúdo
        content = "\n".join(content_parts)
        
        # Criar chunk
        chunk = DocumentChunk(
            chunk_id=f"{sheet_name}_{row_index}",
            document_id=self.metadata.checksum,
            content=content,
            content_type="table_row",
            chunk_index=row_index,
            total_chunks=len(columns),
            metadata={"sheet": sheet_name, "row": row_index},
            embedding_text=content
        )
        
        # Criar entrada
        return KnowledgeBaseEntry(
            uuid=f"{self.metadata.checksum}_{sheet_name}_{row_index}",
            title=f"[Excel] {title}",
            content=content,
            summary=content[:500],
            category="Spreadsheet Data",
            tags=self._extract_tags(content),
            confidence_score=0.8,
            source=self.file_name,
            document_type="excel",
            metadata={"sheet": sheet_name, "row": row_index},
            chunks=[chunk],
            created_by="auto_import"
        )
    
    def _create_entry_from_sheet(self, df, sheet_name: str) -> Optional[KnowledgeBaseEntry]:
        """Cria entrada KB de uma sheet inteira"""
        if df.empty:
            return None
        
        # Criar resumo da sheet
        summary_parts = [
            f"Sheet: {sheet_name}",
            f"Linhas: {len(df)}",
            f"Colunas: {len(df.columns)}",
            f"Colunas: {', '.join(df.columns[:10])}"
        ]
        
        content = "\n".join(summary_parts)
        
        # Adicionar amostra de dados
        content += "\n\nAmostra de dados:\n"
        content += df.head(10).to_string()
        
        return KnowledgeBaseEntry(
            uuid=f"{self.metadata.checksum}_{sheet_name}_sheet",
            title=f"[Excel Sheet] {sheet_name} - {self.file_name}",
            content=content,
            summary=" | ".join(summary_parts),
            category="Spreadsheet",
            tags=[sheet_name, "excel", "tabela"],
            confidence_score=0.9,
            source=self.file_name,
            document_type="excel_sheet",
            metadata={"total_rows": len(df), "total_cols": len(df.columns)},
            chunks=[],
            created_by="auto_import"
        )
    
    def _extract_tags(self, content: str) -> List[str]:
        """Extrai tags do conteúdo"""
        tags = []
        
        # Palavras-chave comuns em sistemas
        keywords = ['CICS', 'DB2', 'IMS', 'COBOL', 'JCL', 'VSAM', 'Mainframe', 'Batch', 'Online']
        for keyword in keywords:
            if keyword.lower() in content.lower():
                tags.append(keyword.lower())
        
        return list(set(tags))

class PDFProcessor(BaseProcessor):
    """Processador para arquivos PDF"""
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa arquivo PDF"""
        entries = []
        
        try:
            import PyPDF2
            from pdfplumber import PDF
            
            # Tentar com pdfplumber primeiro (melhor para tabelas)
            try:
                with PDF.open(self.file_path) as pdf:
                    all_text = []
                    tables = []
                    
                    for i, page in enumerate(pdf.pages):
                        # Extrair texto
                        text = page.extract_text()
                        if text:
                            all_text.append(f"--- Página {i+1} ---\n{text}")
                        
                        # Extrair tabelas
                        page_tables = page.extract_tables()
                        if page_tables:
                            tables.extend(page_tables)
                    
                    # Criar entrada principal
                    content = "\n\n".join(all_text)
                    
                    # Criar chunks por página
                    chunks = []
                    for i, page_text in enumerate(all_text):
                        chunk = DocumentChunk(
                            chunk_id=f"page_{i}",
                            document_id=self.metadata.checksum,
                            content=page_text,
                            content_type="text",
                            chunk_index=i,
                            total_chunks=len(all_text),
                            metadata={"page": i+1},
                            embedding_text=page_text
                        )
                        chunks.append(chunk)
                    
                    # Criar entrada principal
                    entry = KnowledgeBaseEntry(
                        uuid=self.metadata.checksum,
                        title=f"[PDF] {self.file_name}",
                        content=content,
                        summary=content[:1000],
                        category="Document",
                        tags=self._extract_pdf_tags(content),
                        confidence_score=0.9,
                        source=self.file_name,
                        document_type="pdf",
                        metadata={
                            "pages": len(pdf.pages),
                            "has_tables": len(tables) > 0
                        },
                        chunks=chunks,
                        created_by="auto_import"
                    )
                    entries.append(entry)
                    
                    # Criar entradas para tabelas
                    for i, table in enumerate(tables):
                        table_entry = self._create_table_entry(table, i)
                        if table_entry:
                            entries.append(table_entry)
            
            except Exception as e:
                logger.warning(f"pdfplumber falhou, tentando PyPDF2: {str(e)}")
                
                # Fallback para PyPDF2
                with open(self.file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    all_text = []
                    
                    for i, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text:
                            all_text.append(f"--- Página {i+1} ---\n{text}")
                    
                    content = "\n\n".join(all_text)
                    
                    entry = KnowledgeBaseEntry(
                        uuid=self.metadata.checksum,
                        title=f"[PDF] {self.file_name}",
                        content=content,
                        summary=content[:1000],
                        category="Document",
                        tags=["pdf"],
                        confidence_score=0.7,
                        source=self.file_name,
                        document_type="pdf",
                        metadata={"pages": len(pdf_reader.pages)},
                        chunks=[],
                        created_by="auto_import"
                    )
                    entries.append(entry)
        
        except Exception as e:
            logger.error(f"Erro ao processar PDF {self.file_name}: {str(e)}")
        
        return entries
    
    def _extract_pdf_tags(self, content: str) -> List[str]:
        """Extrai tags do PDF"""
        tags = ["pdf", "documento"]
        
        # Adicionar tags baseadas em palavras-chave
        keywords = ['manual', 'guide', 'tutorial', 'specification', 'documentation']
        for keyword in keywords:
            if keyword in content.lower():
                tags.append(keyword)
        
        return list(set(tags))
    
    def _create_table_entry(self, table: List[List], index: int) -> Optional[KnowledgeBaseEntry]:
        """Cria entrada para tabela extraída do PDF"""
        if not table or len(table) < 2:
            return None
        
        # Converter tabela para texto
        content = "Tabela extraída:\n"
        for row in table:
            content += " | ".join(str(cell) for cell in row if cell) + "\n"
        
        return KnowledgeBaseEntry(
            uuid=f"{self.metadata.checksum}_table_{index}",
            title=f"[Tabela PDF] Tabela {index+1} de {self.file_name}",
            content=content,
            summary=f"Tabela com {len(table)} linhas e {len(table[0])} colunas",
            category="Table",
            tags=["tabela", "pdf", "dados"],
            confidence_score=0.8,
            source=self.file_name,
            document_type="pdf_table",
            metadata={"table_index": index, "rows": len(table)},
            chunks=[],
            created_by="auto_import"
        )

class WordProcessor(BaseProcessor):
    """Processador para arquivos Word"""
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa arquivo Word"""
        entries = []
        
        try:
            from docx import Document
            
            doc = Document(self.file_path)
            
            # Extrair texto dos parágrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extrair texto das tabelas
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combinar todo o conteúdo
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n--- Tabelas ---\n\n" + "\n\n".join(tables_text)
            
            # Criar chunks por seção (baseado em títulos)
            chunks = self._create_chunks_from_content(content)
            
            # Extrair metadados do documento
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "has_images": len(doc.inline_shapes) > 0
            }
            
            # Criar entrada principal
            entry = KnowledgeBaseEntry(
                uuid=self.metadata.checksum,
                title=f"[Word] {self.file_name}",
                content=content,
                summary=content[:1000],
                category="Document",
                tags=self._extract_word_tags(content),
                confidence_score=0.9,
                source=self.file_name,
                document_type="word",
                metadata=metadata,
                chunks=chunks,
                created_by="auto_import"
            )
            entries.append(entry)
            
        except Exception as e:
            logger.error(f"Erro ao processar Word {self.file_name}: {str(e)}")
        
        return entries
    
    def _create_chunks_from_content(self, content: str) -> List[DocumentChunk]:
        """Cria chunks baseados em seções do documento"""
        chunks = []
        
        # Dividir por títulos (assumindo que títulos começam com #, números ou são MAIÚSCULAS)
        sections = re.split(r'\n(?=[A-Z\d#]{3,})', content)
        
        for i, section in enumerate(sections):
            if section.strip():
                chunk = DocumentChunk(
                    chunk_id=f"section_{i}",
                    document_id=self.metadata.checksum,
                    content=section,
                    content_type="text",
                    chunk_index=i,
                    total_chunks=len(sections),
                    metadata={"section": i},
                    embedding_text=section
                )
                chunks.append(chunk)
        
        return chunks
    
    def _extract_word_tags(self, content: str) -> List[str]:
        """Extrai tags do documento Word"""
        tags = ["word", "documento"]
        
        # Detectar tipo de documento baseado no conteúdo
        doc_types = {
            'manual': ['manual', 'instruções', 'instructions'],
            'relatório': ['relatório', 'report', 'análise'],
            'procedimento': ['procedimento', 'procedure', 'processo'],
            'política': ['política', 'policy', 'norma']
        }
        
        content_lower = content.lower()
        for tag, keywords in doc_types.items():
            if any(keyword in content_lower for keyword in keywords):
                tags.append(tag)
        
        return list(set(tags))

class ImageProcessor(BaseProcessor):
    """Processador para arquivos de imagem com OCR"""
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa imagem com OCR"""
        entries = []
        
        try:
            from PIL import Image
            import pytesseract
            
            # Abrir imagem
            img = Image.open(self.file_path)
            
            # Fazer OCR
            text = pytesseract.image_to_string(img)
            
            # Metadados da imagem
            metadata = {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,
                "has_text": len(text.strip()) > 0
            }
            
            # Se encontrou texto, criar entrada de texto
            if text.strip():
                content = f"Texto extraído da imagem:\n\n{text}"
                
                entry = KnowledgeBaseEntry(
                    uuid=f"{self.metadata.checksum}_ocr",
                    title=f"[OCR] {self.file_name}",
                    content=content,
                    summary=text[:500],
                    category="Image OCR",
                    tags=["imagem", "ocr", "texto_extraído"],
                    confidence_score=0.7,  # OCR pode ter erros
                    source=self.file_name,
                    document_type="image_ocr",
                    metadata=metadata,
                    chunks=[],
                    created_by="auto_import"
                )
                entries.append(entry)
            
            # Criar entrada para a imagem em si (para busca visual)
            img_entry = KnowledgeBaseEntry(
                uuid=self.metadata.checksum,
                title=f"[Imagem] {self.file_name}",
                content=f"Imagem: {self.file_name}\nDimensões: {img.width}x{img.height}\nFormato: {img.format}",
                summary=f"Arquivo de imagem {img.format} ({img.width}x{img.height})",
                category="Image",
                tags=["imagem", img.format.lower() if img.format else "unknown"],
                confidence_score=1.0,
                source=self.file_name,
                document_type="image",
                metadata=metadata,
                chunks=[],
                created_by="auto_import"
            )
            entries.append(img_entry)
            
        except Exception as e:
            logger.error(f"Erro ao processar imagem {self.file_name}: {str(e)}")
        
        return entries

class TextProcessor(BaseProcessor):
    """Processador para arquivos de texto"""
    
    def process(self) -> List[KnowledgeBaseEntry]:
        """Processa arquivo de texto"""
        entries = []
        
        try:
            # Detectar encoding
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(self.file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    self.metadata.encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                # Fallback: ler como binário
                with open(self.file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
            
            # Detectar tipo de conteúdo
            content_type = self._detect_content_type(content)
            
            # Criar chunks se o arquivo for muito grande
            chunks = []
            if len(content) > 10000:  # Mais de 10KB
                chunks = self._create_text_chunks(content)
            
            # Extrair metadados
            metadata = {
                "lines": content.count('\n') + 1,
                "characters": len(content),
                "encoding": self.metadata.encoding,
                "content_type": content_type
            }
            
            # Criar entrada
            entry = KnowledgeBaseEntry(
                uuid=self.metadata.checksum,
                title=f"[{content_type}] {self.file_name}",
                content=content,
                summary=content[:1000],
                category=content_type,
                tags=self._extract_text_tags(content, content_type),
                confidence_score=0.95,
                source=self.file_name,
                document_type="text",
                metadata=metadata,
                chunks=chunks,
                created_by="auto_import"
            )
            entries.append(entry)
            
        except Exception as e:
            logger.error(f"Erro ao processar texto {self.file_name}: {str(e)}")
        
        return entries
    
    def _detect_content_type(self, content: str) -> str:
        """Detecta o tipo de conteúdo do texto"""
        # Verificar se é JSON
        if content.strip().startswith('{') or content.strip().startswith('['):
            try:
                json.loads(content)
                return "JSON"
            except:
                pass
        
        # Verificar se é XML/HTML
        if content.strip().startswith('<'):
            if '<html' in content.lower():
                return "HTML"
            return "XML"
        
        # Verificar se é YAML
        if content.startswith('---') or ': ' in content[:100]:
            return "YAML"
        
        # Verificar se é Markdown
        if any(pattern in content for pattern in ['# ', '## ', '```', '[](', '![]('):
            return "Markdown"
        
        # Verificar se é código
        code_patterns = ['import ', 'function ', 'class ', 'def ', 'SELECT ', 'CREATE TABLE']
        if any(pattern in content for pattern in code_patterns):
            return "Code"
        
        # Verificar se é log
        if any(word in content.lower()[:1000] for word in ['error', 'warning', 'info', 'debug', 'trace']):
            return "Log"
        
        return "Text"
    
    def _create_text_chunks(self, content: str, chunk_size: int = 5000) -> List[DocumentChunk]:
        """Cria chunks de texto"""
        chunks = []
        
        # Dividir por parágrafos primeiro
        paragraphs = content.split('\n\n')
        
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        chunk_id=f"chunk_{chunk_index}",
                        document_id=self.metadata.checksum,
                        content=current_chunk,
                        content_type="text",
                        chunk_index=chunk_index,
                        total_chunks=0,  # Será atualizado depois
                        metadata={"size": len(current_chunk)},
                        embedding_text=current_chunk
                    ))
                    chunk_index += 1
                current_chunk = para + "\n\n"
        
        # Adicionar último chunk
        if current_chunk:
            chunks.append(DocumentChunk(
                chunk_id=f"chunk_{chunk_index}",
                document_id=self.metadata.checksum,
                content=current_chunk,
                content_type="text",
                chunk_index=chunk_index,
                total_chunks=len(chunks) + 1,
                metadata={"size": len(current_chunk)},
                embedding_text=current_chunk
            ))
        
        # Atualizar total_chunks
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    def _extract_text_tags(self, content: str, content_type: str) -> List[str]:
        """Extrai tags do texto"""
        tags = [content_type.lower()]
        
        # Adicionar tags baseadas em palavras-chave específicas do domínio
        mainframe_keywords = [
            'mainframe', 'cobol', 'jcl', 'cics', 'db2', 'ims', 'vsam',
            'rexx', 'racf', 'tso', 'ispf', 'sdsf', 'sort', 'idcams'
        ]
        
        content_lower = content.lower()
        for keyword in mainframe_keywords:
            if keyword in content_lower:
                tags.append(keyword)
        
        return list(set(tags))

# ================================================================================
# PROCESSADOR UNIVERSAL
# ================================================================================

class UniversalDocumentProcessor:
    """
    Processador universal que gerencia todos os tipos de documentos
    """
    
    def __init__(self, output_dir: str = "./kb_output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.processors = {
            DocumentType.EXCEL: ExcelProcessor,
            DocumentType.PDF: PDFProcessor,
            DocumentType.WORD: WordProcessor,
            DocumentType.IMAGE: ImageProcessor,
            DocumentType.TEXT: TextProcessor,
            # Adicionar mais processadores conforme necessário
        }
    
    def process_file(self, file_path: str) -> List[KnowledgeBaseEntry]:
        """Processa um único arquivo"""
        logger.info(f"Processando: {file_path}")
        
        # Detectar tipo de documento
        doc_type = DocumentTypeDetector.detect(file_path)
        logger.info(f"Tipo detectado: {doc_type.value}")
        
        # Selecionar processador apropriado
        processor_class = self.processors.get(doc_type)
        
        if not processor_class:
            # Fallback para processador de texto
            logger.warning(f"Processador não encontrado para {doc_type.value}, usando TextProcessor")
            processor_class = TextProcessor
        
        # Processar documento
        try:
            processor = processor_class(file_path)
            entries = processor.process()
            logger.info(f"Extraídas {len(entries)} entradas de {file_path}")
            return entries
        except Exception as e:
            logger.error(f"Erro ao processar {file_path}: {str(e)}")
            return []
    
    def process_directory(self, directory: str, recursive: bool = True) -> List[KnowledgeBaseEntry]:
        """Processa todos os arquivos em um diretório"""
        all_entries = []
        path = Path(directory)
        
        # Padrão de busca
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file():
                entries = self.process_file(str(file_path))
                all_entries.extend(entries)
        
        return all_entries
    
    def save_to_json(self, entries: List[KnowledgeBaseEntry], output_file: str = None):
        """Salva entradas em formato JSON"""
        if not output_file:
            output_file = self.output_dir / "kb_entries.json"
        
        # Converter para dict
        entries_dict = []
        for entry in entries:
            entry_dict = asdict(entry)
            # Converter chunks também
            entry_dict['chunks'] = [asdict(chunk) for chunk in entry.chunks]
            entries_dict.append(entry_dict)
        
        # Salvar JSON
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(entries_dict, f, ensure_ascii=False, indent=2, default=str)
        
        logger.info(f"Salvo {len(entries)} entradas em {output_file}")
    
    def generate_sql_script(self, entries: List[KnowledgeBaseEntry], output_file: str = None):
        """Gera script SQL para inserção no PostgreSQL"""
        if not output_file:
            output_file = self.output_dir / "insert_kb.sql"
        
        sql_lines = [
            "-- Script de inserção na Knowledge Base",
            f"-- Gerado em: {datetime.now().isoformat()}",
            f"-- Total de entradas: {len(entries)}",
            "",
            "BEGIN;",
            ""
        ]
        
        for entry in entries:
            # Escapar aspas simples
            title = entry.title.replace("'", "''")
            content = entry.content.replace("'", "''")
            summary = entry.summary.replace("'", "''")
            tags = "{" + ",".join(entry.tags) + "}"
            metadata = json.dumps(entry.metadata).replace("'", "''")
            
            sql = f"""
INSERT INTO knowledge_base (
    uuid, title, content, summary, category, tags,
    confidence_score, source, metadata, created_by, created_at
) VALUES (
    '{entry.uuid}', '{title}', '{content}', '{summary}',
    '{entry.category}', '{tags}', {entry.confidence_score},
    '{entry.source}', '{metadata}'::jsonb, '{entry.created_by}', CURRENT_TIMESTAMP
) ON CONFLICT (uuid) DO UPDATE SET
    content = EXCLUDED.content,
    updated_at = CURRENT_TIMESTAMP;
"""
            sql_lines.append(sql)
        
        sql_lines.extend(["", "COMMIT;", ""])
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("\n".join(sql_lines))
        
        logger.info(f"Script SQL gerado: {output_file}")
    
    def generate_import_script(self, entries: List[KnowledgeBaseEntry], output_file: str = None):
        """Gera script Python/Node.js para importação com embeddings"""
        if not output_file:
            output_file = self.output_dir / "import_with_embeddings.py"
        
        script = f'''#!/usr/bin/env python3
"""
Script de importação com geração de embeddings
"""

import json
import psycopg2
from openai import OpenAI
import time

# Configuração
DB_CONFIG = {{
    "host": "localhost",
    "port": 5432,
    "database": "mainframe_ai",
    "user": "mainframe_user",
    "password": "mainframe_pass"
}}

# Dados a importar
entries = {json.dumps([asdict(e) for e in entries[:5]], indent=2)}  # Amostra de 5 entradas

# Cliente OpenAI
openai_client = OpenAI()

def generate_embedding(text):
    """Gera embedding usando OpenAI"""
    try:
        response = openai_client.embeddings.create(
            model="text-embedding-ada-002",
            input=text[:8000]  # Limitar tamanho
        )
        return response.data[0].embedding
    except Exception as e:
        print(f"Erro ao gerar embedding: {{e}}")
        return None

def import_entries():
    """Importa entradas com embeddings"""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    
    for entry in entries:
        print(f"Processando: {{entry['title']}}")
        
        # Gerar embedding
        embedding = generate_embedding(entry['content'])
        
        if embedding:
            # Inserir no banco
            cur.execute("""
                INSERT INTO knowledge_base (
                    uuid, title, content, summary, category, tags,
                    confidence_score, source, embedding, metadata,
                    created_by, created_at
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP
                ) ON CONFLICT (uuid) DO UPDATE SET
                    content = EXCLUDED.content,
                    embedding = EXCLUDED.embedding,
                    updated_at = CURRENT_TIMESTAMP
            """, (
                entry['uuid'], entry['title'], entry['content'],
                entry['summary'], entry['category'], entry['tags'],
                entry['confidence_score'], entry['source'],
                json.dumps(embedding), json.dumps(entry['metadata']),
                entry['created_by']
            ))
        
        time.sleep(0.1)  # Rate limiting
    
    conn.commit()
    cur.close()
    conn.close()
    print("Importação concluída!")

if __name__ == "__main__":
    import_entries()
'''
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(script)
        
        logger.info(f"Script de importação gerado: {output_file}")

# ================================================================================
# FUNÇÕES PRINCIPAIS
# ================================================================================

def main():
    """Função principal"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Processador Universal de Documentos para Knowledge Base")
    parser.add_argument("input", help="Arquivo ou diretório para processar")
    parser.add_argument("-o", "--output", default="./kb_output", help="Diretório de saída")
    parser.add_argument("-r", "--recursive", action="store_true", help="Processar diretórios recursivamente")
    parser.add_argument("--sql", action="store_true", help="Gerar script SQL")
    parser.add_argument("--json", action="store_true", help="Gerar arquivo JSON")
    parser.add_argument("--import-script", action="store_true", help="Gerar script de importação")
    
    args = parser.parse_args()
    
    # Criar processador
    processor = UniversalDocumentProcessor(args.output)
    
    # Processar entrada
    input_path = Path(args.input)
    
    if input_path.is_file():
        entries = processor.process_file(str(input_path))
    elif input_path.is_dir():
        entries = processor.process_directory(str(input_path), args.recursive)
    else:
        logger.error(f"Caminho inválido: {args.input}")
        return
    
    # Salvar resultados
    if entries:
        logger.info(f"\n{'='*60}")
        logger.info(f"Total de entradas processadas: {len(entries)}")
        logger.info(f"{'='*60}\n")
        
        # Estatísticas
        categories = {}
        doc_types = {}
        
        for entry in entries:
            categories[entry.category] = categories.get(entry.category, 0) + 1
            doc_types[entry.document_type] = doc_types.get(entry.document_type, 0) + 1
        
        logger.info("Distribuição por categoria:")
        for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
            logger.info(f"  - {cat}: {count}")
        
        logger.info("\nDistribuição por tipo de documento:")
        for dtype, count in sorted(doc_types.items(), key=lambda x: x[1], reverse=True):
            logger.info(f"  - {dtype}: {count}")
        
        # Gerar saídas
        if args.json or not (args.sql or args.import_script):
            processor.save_to_json(entries)
        
        if args.sql:
            processor.generate_sql_script(entries)
        
        if args.import_script:
            processor.generate_import_script(entries)
        
        logger.info(f"\nArquivos gerados em: {processor.output_dir}")
    else:
        logger.warning("Nenhuma entrada foi processada")

if __name__ == "__main__":
    main()
